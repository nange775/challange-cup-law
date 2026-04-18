from typing import Any, Dict, List, Optional
from docx import Document
from .base import BaseParser
from .registry import ParserRegistry

class WordParser(BaseParser):
    supported_extensions = ['.docx']
    supported_mimes = [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    ]

    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        doc = Document(file_path)

        # 提取段落
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # 提取表格
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)

        full_content = "\n".join(paragraphs)

        # 提取实体
        entities = self.extract_entities(full_content)

        return {
            "content": full_content,
            "structured_data": {
                "paragraphs": paragraphs,
                "tables": tables_data,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables_data)
            },
            "metadata": {
                "title": doc.core_properties.title or paragraphs[0] if paragraphs else "",
                "author": doc.core_properties.author,
                "created": doc.core_properties.created
            },
            "entities": entities,
            "type": "word_document"
        }

    def extract_entities(self, content: str) -> Dict[str, List[str]]:
        import re
        # 复用PDF解析器的实体提取逻辑
        names = re.findall(r'[姓名][:：]\s*([\u4e00-\u9fa5]{2,4})', content)
        amounts = re.findall(r'(\d+\.?\d*)\s*[元万元]', content)
        dates = re.findall(r'\d{4}[-/年]\d{1,2}[-/月]\d{1,2}[日]?', content)

        return {
            "names": list(set(names)),
            "amounts": list(set(amounts)),
            "dates": list(set(dates))
        }

ParserRegistry.register(WordParser())